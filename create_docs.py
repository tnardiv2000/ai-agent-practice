from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('AI Data Analyzer Pro', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('A Smart Data Analysis Application')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.italic = True

doc.add_paragraph()  # Blank line

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc_items = [
    '1. Overview',
    '2. Features',
    '3. Requirements',
    '4. Installation',
    '5. Configuration',
    '6. How to Use',
    '7. Code Walkthrough',
    '8. Project Structure',
    '9. Troubleshooting'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Overview Section
doc.add_heading('Overview', 1)
doc.add_paragraph(
    'AI Data Analyzer Pro is a Streamlit web application that allows you to:'
)
overview_items = [
    'Upload Excel or CSV files',
    'Perform smart data aggregations with intelligent function selection',
    'Get verified, accurate results matching your pivot tables',
    'Ask AI to analyze patterns and provide insights'
]
for item in overview_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'Key Innovation: Results are verified through manual calculations BEFORE AI analysis, ensuring accuracy and preventing hallucination.'
)

doc.add_page_break()

# Features Section
doc.add_heading('Features', 1)

doc.add_heading('📊 Data Preview & Statistics', 2)
doc.add_paragraph(
    'View raw data with customizable row display\nQuick statistics dashboard (row count, column count, missing values, file size)\nDetailed column information including data types and unique value counts',
    style='List Bullet'
)

doc.add_heading('🧮 Manual Data Calculations', 2)
doc.add_paragraph('All calculations happen first, verified data goes to AI:')
calc_items = [
    'Column Statistics - Detailed stats (min, max, mean, median, std dev) for any column',
    'Filter & View - Filter data by column values and view complete rows',
    'Group & Aggregate - Smart aggregation with intelligent function selection',
    'Custom View - Select specific columns to compare'
]
for item in calc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('🤖 AI-Powered Analysis', 2)
ai_items = [
    'Ask questions about your data',
    'AI references verified aggregation results',
    'Temperature control for precision vs. creativity',
    'Timeout settings for long-running queries',
    'Show/hide analysis steps for transparency'
]
for item in ai_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Requirements Section
doc.add_heading('Requirements', 1)

doc.add_heading('Software', 2)
software = ['Python 3.8+', 'Ollama (for local AI/LLM)', 'Git (optional, for version control)']
for soft in software:
    doc.add_paragraph(soft, style='List Bullet')

doc.add_heading('Python Libraries', 2)
doc.add_paragraph('streamlit>=1.28.0')
doc.add_paragraph('pandas>=1.5.0')
doc.add_paragraph('openpyxl>=3.10.0')
doc.add_paragraph('xlrd>=2.0.1')
doc.add_paragraph('python-dotenv>=1.0.0')
doc.add_paragraph('requests>=2.31.0')

doc.add_page_break()

# Installation Section
doc.add_heading('Installation', 1)

doc.add_heading('Step 1: Clone or Create Project', 2)
doc.add_paragraph('mkdir ai-agent-practice')
doc.add_paragraph('cd ai-agent-practice')

doc.add_heading('Step 2: Create Virtual Environment', 2)
doc.add_paragraph('Windows: python -m venv venv')
doc.add_paragraph('        venv\\Scripts\\activate')
doc.add_paragraph()
doc.add_paragraph('Mac/Linux: python3 -m venv venv')
doc.add_paragraph('          source venv/bin/activate')

doc.add_heading('Step 3: Install Dependencies', 2)
doc.add_paragraph('pip install -r requirements.txt')

doc.add_heading('Step 4: Install Ollama', 2)
doc.add_paragraph('Download and install from: https://ollama.ai')
doc.add_paragraph('Then pull the LLM model: ollama pull llama2')

doc.add_page_break()

# Configuration Section
doc.add_heading('Configuration', 1)

doc.add_heading('Step 1: Create .env File', 2)
doc.add_paragraph('Create a file named .env in your project root:')
doc.add_paragraph('OLLAMA_API=http://localhost:11434/api/generate')

doc.add_heading('Step 2: Verify Ollama is Running', 2)
doc.add_paragraph('Open a terminal and run: ollama serve')
doc.add_paragraph('You should see: Ollama is running on http://localhost:11434')
doc.add_paragraph('Keep this terminal open while using the app.')

doc.add_page_break()

# How to Use Section
doc.add_heading('How to Use', 1)

doc.add_heading('Starting the App', 2)
doc.add_paragraph('Option 1 - Direct start: streamlit run main.py')
doc.add_paragraph('Option 2 - Using restart script: restart.bat')
doc.add_paragraph('App opens at: http://localhost:8501')

doc.add_heading('Workflow', 2)

doc.add_heading('1. Upload Your Data', 3)
doc.add_paragraph('Click Browse files')
doc.add_paragraph('Select Excel (.xlsx, .xls) or CSV file')
doc.add_paragraph('App loads and displays initial statistics')

doc.add_heading('2. Explore Data', 3)
doc.add_paragraph('Data Preview tab - See raw data', style='List Bullet')
doc.add_paragraph('Data Stats tab - Quick statistics and numeric summaries', style='List Bullet')
doc.add_paragraph('Detailed Info tab - Column types, unique values, missing data', style='List Bullet')

doc.add_heading('3. Manual Calculations (Most Important!)', 3)
doc.add_paragraph('Use ONE of these tabs to verify exact numbers:')
doc.add_paragraph('Option A: Column Statistics', style='List Bullet')
doc.add_paragraph('Option B: Filter & View', style='List Bullet')
doc.add_paragraph('Option C: Group & Aggregate (RECOMMENDED)', style='List Bullet')
doc.add_paragraph('Option D: Custom View', style='List Bullet')

doc.add_heading('4. AI Analysis (After Verification)', 3)
doc.add_paragraph('Once you have verified your numbers:')
doc.add_paragraph('Type a question in the text area', style='List Bullet')
doc.add_paragraph('Example: Why does APAC Japan have higher max spend than India?', style='List Bullet')
doc.add_paragraph('AI references your verified numbers', style='List Bullet')
doc.add_paragraph('No hallucination possible!', style='List Bullet')

doc.add_page_break()

# Project Structure
doc.add_heading('Project Structure', 1)
doc.add_paragraph('''
ai-agent-practice/
├── venv/                          # Virtual environment (ignore in git)
│   └── [Python packages]
│
├── main.py                        # Main application file (550+ lines)
├── requirements.txt               # Python dependencies
├── .env                          # Environment variables (local, not in git)
├── .env.example                  # Example env file (for reference)
├── .gitignore                    # Git ignore rules
├── restart.bat                   # Quick restart script (Windows)
└── README.md                     # Documentation
''')

doc.add_page_break()

# Troubleshooting
doc.add_heading('Troubleshooting', 1)

doc.add_heading('"Cannot connect to Ollama"', 2)
doc.add_paragraph('Start Ollama: ollama serve in separate terminal', style='List Bullet')
doc.add_paragraph('Verify running at: http://localhost:11434', style='List Bullet')

doc.add_heading('"Streamlit won\'t stop"', 2)
doc.add_paragraph('Use PowerShell as Admin:', style='List Bullet')
doc.add_paragraph('Get-Process -Id (Get-NetTCPConnection -LocalPort 8501).OwningProcess | Stop-Process -Force')

doc.add_heading('"ModuleNotFoundError"', 2)
doc.add_paragraph('Verify virtual environment activated', style='List Bullet')
doc.add_paragraph('Run: pip install -r requirements.txt', style='List Bullet')

doc.add_heading('"AI giving wrong numbers"', 2)
doc.add_paragraph('Use Group & Aggregate tab FIRST to verify', style='List Bullet')
doc.add_paragraph('Compare results with Excel pivot table', style='List Bullet')
doc.add_paragraph('Only then ask AI follow-up questions', style='List Bullet')

doc.add_page_break()

# Technologies
doc.add_heading('What Was Needed to Create This', 1)

doc.add_heading('Technologies', 2)
tech_items = [
    'Streamlit - Web framework for data apps',
    'Pandas - Data manipulation and analysis',
    'Ollama - Local LLM running environment',
    'Llama2 - AI model for analysis',
    'Python - Programming language',
    'Requests - HTTP library for API calls',
    'python-dotenv - Environment variable management'
]
for tech in tech_items:
    doc.add_paragraph(tech, style='List Bullet')

doc.add_heading('Key Concepts Used', 2)
concepts = [
    'Smart type detection for columns',
    'Function recommendation system',
    'Session state management',
    'API integration with local LLM',
    'Data aggregation and grouping',
    'File I/O handling',
    'Error handling and user feedback'
]
for concept in concepts:
    doc.add_paragraph(concept, style='List Bullet')

doc.add_page_break()

# Getting Started
doc.add_heading('Getting Started', 1)
steps = [
    'Install: pip install -r requirements.txt',
    'Configure: Create .env file with Ollama URL',
    'Start Ollama: ollama serve (in separate terminal)',
    'Run app: streamlit run main.py',
    'Upload data: Use file uploader',
    'Verify numbers: Use Group & Aggregate tab',
    'Ask AI: Use AI analysis section'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

# Save document
doc.save('AI_Data_Analyzer_Pro_Documentation.docx')
print("✅ Word document created: AI_Data_Analyzer_Pro_Documentation.docx")